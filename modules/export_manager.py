import os
import json
from datetime import datetime
from typing import Dict, Optional
import markdown2
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from docx import Document
import logging

class ExportManager:
    def __init__(self, config):
        self.config = config
        self.output_dir = config.output_dir
        os.makedirs(self.output_dir, exist_ok=True)
    
    def save_article(self, article_data: Dict, filename: str, 
                    export_format: str, include_bibliography: bool = True,
                    include_metadata: bool = True) -> str:
        """Główna metoda eksportu artykułu"""
        
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        base_filename = f"{filename}_{timestamp}"
        
        try:
            if export_format == "html":
                filepath = self._save_html(article_data, base_filename, 
                                         include_bibliography, include_metadata)
            elif export_format == "markdown":
                filepath = self._save_markdown(article_data, base_filename,
                                             include_bibliography, include_metadata)
            elif export_format == "pdf":
                filepath = self._save_pdf(article_data, base_filename,
                                        include_bibliography, include_metadata)
            elif export_format == "docx":
                filepath = self._save_docx(article_data, base_filename,
                                         include_bibliography, include_metadata)
            elif export_format == "txt":
                filepath = self._save_txt(article_data, base_filename,
                                        include_bibliography, include_metadata)
            else:
                raise ValueError(f"Nieobsługiwany format eksportu: {export_format}")
            
            return filepath
            
        except Exception as e:
            logging.error(f"Błąd eksportu do {export_format}: {e}")
            raise
    
    def _save_html(self, article_data: Dict, filename: str,
                   include_bibliography: bool, include_metadata: bool) -> str:
        """Eksport do HTML"""
        filepath = os.path.join(self.output_dir, f"{filename}.html")
        
        html_content = f"""<!DOCTYPE html>
<html lang="pl">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{article_data['metadata']['topic']}</title>
    <style>
        body {{ font-family: Arial, sans-serif; line-height: 1.6; margin: 40px; }}
        h1 {{ color: #333; border-bottom: 2px solid #ccc; }}
        h2 {{ color: #666; }}
        .metadata {{ background: #f4f4f4; padding: 15px; margin: 20px 0; border-radius: 5px; }}
        .sources {{ margin-top: 30px; }}
        .source {{ margin: 10px 0; padding: 10px; background: #f9f9f9; border-left: 3px solid #007cba; }}
    </style>
</head>
<body>
    <h1>{article_data['metadata']['topic']}</h1>
"""
        
        if include_metadata:
            html_content += f"""
    <div class="metadata">
        <h3>Metadane</h3>
        <p><strong>Typ:</strong> {article_data['metadata']['type']}</p>
        <p><strong>Długość:</strong> {article_data['metadata']['length']}</p>
        <p><strong>Język:</strong> {article_data['metadata']['language']}</p>
        <p><strong>Model:</strong> {article_data['metadata']['provider']} - {article_data['metadata']['model']}</p>
        <p><strong>Liczba słów:</strong> {article_data.get('word_count', 'N/A')}</p>
        <p><strong>Ocena jakości:</strong> {article_data['metadata']['quality_score']:.2f}/1.0</p>
    </div>
"""
        
        # Konwertuj markdown na HTML
        html_content += markdown2.markdown(article_data['content'])
        
        if include_bibliography and 'sources' in article_data:
            html_content += """
    <div class="sources">
        <h2>Bibliografia</h2>
"""
            for i, source in enumerate(article_data['sources'], 1):
                html_content += f"""
        <div class="source">
            <strong>{i}. {source.get('title', 'Bez tytułu')}</strong><br>
            <em>Typ:</em> {source.get('type', 'nieznany')}<br>
            <em>URL:</em> <a href="{source.get('url', '#')}" target="_blank">{source.get('url', 'Brak URL')}</a><br>
            <em>Wiarygodność:</em> {source.get('credibility', 0):.1f}/1.0
        </div>
"""
            html_content += "    </div>"
        
        html_content += """
</body>
</html>
"""
        
        with open(filepath, 'w', encoding='utf-8') as f:
            f.write(html_content)
        
        return filepath
    
    def _save_markdown(self, article_data: Dict, filename: str,
                      include_bibliography: bool, include_metadata: bool) -> str:
        """Eksport do Markdown"""
        filepath = os.path.join(self.output_dir, f"{filename}.md")
        
        content = f"# {article_data['metadata']['topic']}\n\n"
        
        if include_metadata:
            content += "## Metadane\n\n"
            content += f"- **Typ:** {article_data['metadata']['type']}\n"
            content += f"- **Długość:** {article_data['metadata']['length']}\n"
            content += f"- **Język:** {article_data['metadata']['language']}\n"
            content += f"- **Model:** {article_data['metadata']['provider']} - {article_data['metadata']['model']}\n"
            content += f"- **Liczba słów:** {article_data.get('word_count', 'N/A')}\n"
            content += f"- **Ocena jakości:** {article_data['metadata']['quality_score']:.2f}/1.0\n\n"
        
        content += article_data['content']
        
        if include_bibliography and 'sources' in article_data:
            content += "\n\n## Bibliografia\n\n"
            for i, source in enumerate(article_data['sources'], 1):
                content += f"{i}. **{source.get('title', 'Bez tytułu')}**\n"
                content += f"   - Typ: {source.get('type', 'nieznany')}\n"
                content += f"   - URL: [{source.get('url', 'Brak URL')}]({source.get('url', '#')})\n"
                content += f"   - Wiarygodność: {source.get('credibility', 0):.1f}/1.0\n\n"
        
        with open(filepath, 'w', encoding='utf-8') as f:
            f.write(content)
        
        return filepath
    
    def _save_pdf(self, article_data: Dict, filename: str,
                  include_bibliography: bool, include_metadata: bool) -> str:
        """Eksport do PDF"""
        filepath = os.path.join(self.output_dir, f"{filename}.pdf")
        
        doc = SimpleDocTemplate(filepath, pagesize=letter)
        styles = getSampleStyleSheet()
        story = []
        
        # Tytuł
        title = Paragraph(article_data['metadata']['topic'], styles['Title'])
        story.append(title)
        story.append(Spacer(1, 12))
        
        # Treść artykułu (uproszczona)
        content_lines = article_data['content'].split('\n')
        for line in content_lines[:50]:  # Limit dla przykładu
            if line.strip():
                para = Paragraph(line, styles['Normal'])
                story.append(para)
                story.append(Spacer(1, 6))
        
        doc.build(story)
        return filepath
    
    def _save_docx(self, article_data: Dict, filename: str,
                   include_bibliography: bool, include_metadata: bool) -> str:
        """Eksport do DOCX"""
        filepath = os.path.join(self.output_dir, f"{filename}.docx")
        
        doc = Document()
        
        # Tytuł
        title = doc.add_heading(article_data['metadata']['topic'], 0)
        
        # Treść artykułu (uproszczona)
        content_lines = article_data['content'].split('\n')
        for line in content_lines:
            if line.strip():
                if line.startswith('#'):
                    level = line.count('#')
                    clean_line = line.lstrip('# ').strip()
                    doc.add_heading(clean_line, level=min(level, 3))
                else:
                    doc.add_paragraph(line)
        
        doc.save(filepath)
        return filepath
    
    def _save_txt(self, article_data: Dict, filename: str,
                  include_bibliography: bool, include_metadata: bool) -> str:
        """Eksport do TXT"""
        filepath = os.path.join(self.output_dir, f"{filename}.txt")
        
        content = f"{article_data['metadata']['topic']}\n"
        content += "=" * len(article_data['metadata']['topic']) + "\n\n"
        content += article_data['content']
        
        with open(filepath, 'w', encoding='utf-8') as f:
            f.write(content)
        
        return filepath

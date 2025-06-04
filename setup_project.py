#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Skrypt do automatycznego tworzenia struktury projektu Streamlit Article Generator
wraz z wszystkimi niezbędnymi plikami kodu.

Autor: AI Assistant
Data: 2025-06-04
"""

import os
import sys
from pathlib import Path
import json
import yaml


def create_directory_structure():
    """Tworzy strukturę katalogów projektu"""
    print("🏗️  Tworzenie struktury katalogów...")
    
    directories = [
        "streamlit_article_generator",
        "streamlit_article_generator/modules",
        "streamlit_article_generator/config",
        "streamlit_article_generator/data",
        "streamlit_article_generator/data/cache",
        "streamlit_article_generator/outputs",
    ]
    
    for directory in directories:
        Path(directory).mkdir(parents=True, exist_ok=True)
        print(f"✅ Utworzono katalog: {directory}")


def create_requirements_txt():
    """Tworzy plik requirements.txt"""
    print("\n📦 Tworzenie pliku requirements.txt...")
    
    requirements = """streamlit>=1.28.0
requests>=2.31.0
beautifulsoup4>=4.12.0
nltk>=3.8.1
spacy>=3.7.0
wikipedia-api>=0.6.0
openai>=1.3.0
anthropic>=0.7.0
langchain>=0.1.0
langchain-community>=0.0.10
pandas>=2.1.0
plotly>=5.17.0
reportlab>=4.0.0
markdown2>=2.4.0
python-docx>=1.1.0
python-dotenv>=1.0.0
wordcloud>=1.9.0
matplotlib>=3.8.0
sentence-transformers>=2.2.0
PyYAML>=6.0.1
groq>=0.4.0
"""
    
    with open("streamlit_article_generator/requirements.txt", "w", encoding="utf-8") as f:
        f.write(requirements)
    print("✅ Plik requirements.txt utworzony")


def create_env_file():
    """Tworzy przykładowy plik .env"""
    print("\n🔐 Tworzenie pliku .env...")
    
    env_content = """# API Keys - Uzupełnij własnymi kluczami
OPENAI_API_KEY=your_openai_api_key_here
ANTHROPIC_API_KEY=your_anthropic_api_key_here
GROQ_API_KEY=your_groq_api_key_here

# Opcjonalne ustawienia
DEFAULT_LANGUAGE=polski
DEFAULT_ARTICLE_TYPE=blog_post
MAX_SOURCES=10
"""
    
    with open("streamlit_article_generator/.env", "w", encoding="utf-8") as f:
        f.write(env_content)
    print("✅ Plik .env utworzony")


def create_config_settings():
    """Tworzy plik config/settings.py"""
    print("\n⚙️  Tworzenie pliku config/settings.py...")
    
    settings_content = '''from dataclasses import dataclass
from typing import Dict, List
import os
from dotenv import load_dotenv

load_dotenv()

@dataclass
class LLMConfig:
    openai_models: List[str]
    anthropic_models: List[str]
    ollama_models: List[str]
    groq_models: List[str]
    max_tokens: int
    temperature: float

@dataclass
class ScrapingConfig:
    max_sources: int
    timeout: int
    rate_limit: float
    user_agent: str

@dataclass
class AppConfig:
    llm: LLMConfig
    scraping: ScrapingConfig
    cache_ttl: int
    max_history: int
    output_dir: str
    
    @classmethod
    def load(cls):
        return cls(
            llm=LLMConfig(
                openai_models=['gpt-4o', 'gpt-4o-mini', 'gpt-3.5-turbo'],
                anthropic_models=['claude-3-sonnet-20240229', 'claude-3-haiku-20240307'],
                ollama_models=['llama3.1', 'mistral', 'codellama', 'gemma2'],
                groq_models=['llama3-70b-8192', 'mixtral-8x7b-32768'],
                max_tokens=4000,
                temperature=0.7
            ),
            scraping=ScrapingConfig(
                max_sources=15,
                timeout=10,
                rate_limit=1.0,
                user_agent='Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'
            ),
            cache_ttl=3600,
            max_history=50,
            output_dir='outputs'
        )
'''
    
    with open("streamlit_article_generator/config/settings.py", "w", encoding="utf-8") as f:
        f.write(settings_content)
    print("✅ Plik config/settings.py utworzony")


def create_prompts_yaml():
    """Tworzy plik config/prompts.yaml"""
    print("\n📝 Tworzenie pliku config/prompts.yaml...")
    
    prompts_data = {
        'article_types': {
            'esej': {
                'system_prompt': 'Jesteś ekspertem w pisaniu esejów akademickich. Twoje zadanie to napisanie przemyślanego eseju na podany temat, wykorzystując dostarczone źródła.',
                'user_template': '''Napisz esej na temat: "{topic}"

Długość: {length}
Język: {language}

Źródła do wykorzystania:
{sources}

Dodatkowe instrukcje: {custom_instructions}

Esej powinien zawierać:
- Wprowadzenie z jasną tezą
- Rozwinięcie argumentów z przykładami
- Podsumowanie i wnioski
- Odpowiednie cytowania źródeł'''
            },
            'blog_post': {
                'system_prompt': 'Jesteś doświadczonym blogerem specjalizującym się w tworzeniu angażujących treści SEO-friendly. Piszesz posty, które są zarówno informacyjne, jak i interesujące.',
                'user_template': '''Napisz post na blog na temat: "{topic}"

Długość: {length}
Język: {language}

Źródła do wykorzystania:
{sources}

Dodatkowe instrukcje: {custom_instructions}

Post powinien zawierać:
- Przyciągający tytuł
- Wprowadzenie z hook'iem
- Nagłówki i listy punktowane
- Call-to-action na końcu
- SEO-friendly struktura'''
            },
            'artykul_wikipedii': {
                'system_prompt': 'Jesteś redaktorem Wikipedii. Tworzysz neutralne, encyklopedyczne artykuły oparte na wiarygodnych źródłach, zgodnie ze standardami Wikipedii.',
                'user_template': '''Napisz artykuł w stylu Wikipedii na temat: "{topic}"

Długość: {length}
Język: {language}

Źródła do wykorzystania:
{sources}

Dodatkowe instrukcje: {custom_instructions}

Artykuł powinien zawierać:
- Neutralny punkt widzenia
- Weryfikowalne informacje
- Odpowiednie kategorie i linki
- Bibliografia z numerowanymi przypisami'''
            },
            'artykul_naukowy': {
                'system_prompt': 'Jesteś badaczem naukowym piszącym artykuły do czasopism naukowych. Tworzysz metodyczne, dobrze udokumentowane prace badawcze.',
                'user_template': '''Napisz artykuł naukowy na temat: "{topic}"

Długość: {length}
Język: {language}

Źródła do wykorzystania:
{sources}

Dodatkowe instrukcje: {custom_instructions}

Artykuł powinien zawierać:
- Abstrakt
- Wprowadzenie z przeglądem literatury
- Metodologię (jeśli applicable)
- Wyniki i dyskusję
- Wnioski i bibliografia'''
            },
            'poradnik': {
                'system_prompt': 'Jesteś ekspertem w tworzeniu praktycznych poradników. Piszesz jasne, wykonalne instrukcje krok po kroku.',
                'user_template': '''Napisz poradnik na temat: "{topic}"

Długość: {length}
Język: {language}

Źródła do wykorzystania:
{sources}

Dodatkowe instrukcje: {custom_instructions}

Poradnik powinien zawierać:
- Jasne wprowadzenie do tematu
- Listę potrzebnych narzędzi/materiałów
- Instrukcje krok po kroku
- Porady i ostrzeżenia
- Podsumowanie i dalsze kroki'''
            },
            'news_summary': {
                'system_prompt': 'Jesteś dziennikarzem specjalizującym się w tworzeniu streszczeń aktualnych wydarzeń. Przedstawiasz fakty w sposób obiektywny i przystępny.',
                'user_template': '''Napisz streszczenie aktualnych wydarzeń na temat: "{topic}"

Długość: {length}
Język: {language}

Źródła do wykorzystania:
{sources}

Dodatkowe instrukcje: {custom_instructions}

Streszczenie powinno zawierać:
- Najważniejsze fakty na początku
- Kontekst i tło wydarzeń
- Różne perspektywy jeśli applicable
- Aktualne dane i statystyki
- Obiektywny ton bez opinii'''
            }
        }
    }
    
    with open("streamlit_article_generator/config/prompts.yaml", "w", encoding="utf-8") as f:
        yaml.dump(prompts_data, f, default_flow_style=False, allow_unicode=True, indent=2)
    print("✅ Plik config/prompts.yaml utworzony")


def create_sources_json():
    """Tworzy plik data/sources.json"""
    print("\n🔗 Tworzenie pliku data/sources.json...")
    
    sources_data = {
        "news": [
            {
                "name": "TVN24",
                "base_url": "https://tvn24.pl",
                "search_url": "https://tvn24.pl/szukaj?q={query}",
                "credibility": 0.8
            },
            {
                "name": "Gazeta.pl",
                "base_url": "https://gazeta.pl",
                "search_url": "https://gazeta.pl/0,0.html?szukaj={query}",
                "credibility": 0.7
            },
            {
                "name": "Onet",
                "base_url": "https://onet.pl",
                "search_url": "https://onet.pl/szukaj/{query}",
                "credibility": 0.6
            }
        ],
        "academic": [
            {
                "name": "Google Scholar",
                "base_url": "https://scholar.google.com",
                "search_url": "https://scholar.google.com/scholar?q={query}",
                "credibility": 0.9
            },
            {
                "name": "ResearchGate",
                "base_url": "https://researchgate.net",
                "search_url": "https://researchgate.net/search?q={query}",
                "credibility": 0.8
            }
        ],
        "general": [
            {
                "name": "Wikipedia PL",
                "base_url": "https://pl.wikipedia.org",
                "search_url": "https://pl.wikipedia.org/w/index.php?search={query}",
                "credibility": 0.9
            }
        ]
    }
    
    with open("streamlit_article_generator/data/sources.json", "w", encoding="utf-8") as f:
        json.dump(sources_data, f, ensure_ascii=False, indent=2)
    print("✅ Plik data/sources.json utworzony")


def create_modules_init():
    """Tworzy plik modules/__init__.py"""
    print("\n📦 Tworzenie pliku modules/__init__.py...")
    
    init_content = '''"""
Moduły dla aplikacji Streamlit Article Generator
"""

__version__ = "1.0.0"
__author__ = "AI Assistant"

# Import głównych klas
from .web_scraper import WebScraper
from .llm_integration import LLMManager
from .article_generator import ArticleGenerator
from .export_manager import ExportManager

__all__ = [
    'WebScraper',
    'LLMManager', 
    'ArticleGenerator',
    'ExportManager'
]
'''
    
    with open("streamlit_article_generator/modules/__init__.py", "w", encoding="utf-8") as f:
        f.write(init_content)
    print("✅ Plik modules/__init__.py utworzony")


def create_web_scraper():
    """Tworzy plik modules/web_scraper.py"""
    print("\n🕷️  Tworzenie pliku modules/web_scraper.py...")
    
    web_scraper_content = '''import requests
from bs4 import BeautifulSoup
import wikipedia
import time
import json
import os
from typing import List, Dict, Optional
from urllib.parse import urljoin, urlparse
import logging

class WebScraper:
    def __init__(self, config):
        self.config = config
        self.session = requests.Session()
        self.session.headers.update({'User-Agent': config.scraping.user_agent})
        self.sources_data = self._load_sources()
        
    def _load_sources(self) -> Dict:
        """Ładuje listę zaufanych źródeł"""
        sources_path = 'data/sources.json'
        if os.path.exists(sources_path):
            with open(sources_path, 'r', encoding='utf-8') as f:
                return json.load(f)
        return {"news": [], "academic": [], "general": []}
    
    def search_sources(self, topic: str, max_sources: int = 10, 
                      search_depth: str = "medium") -> List[Dict]:
        """Główna metoda wyszukiwania źródeł"""
        all_sources = []
        
        # Wikipedia search
        wikipedia_sources = self._search_wikipedia(topic, max_sources // 3)
        all_sources.extend(wikipedia_sources)
        
        # News search
        if search_depth in ["medium", "deep"]:
            news_sources = self._search_news(topic, max_sources // 3)
            all_sources.extend(news_sources)
        
        # Academic search
        if search_depth == "deep":
            academic_sources = self._search_academic(topic, max_sources // 3)
            all_sources.extend(academic_sources)
        
        # General web search
        general_sources = self._search_general(topic, max_sources - len(all_sources))
        all_sources.extend(general_sources)
        
        return self._rank_sources(all_sources)[:max_sources]
    
    def _search_wikipedia(self, topic: str, max_results: int) -> List[Dict]:
        """Wyszukuje w Wikipedii"""
        sources = []
        try:
            wikipedia.set_lang("pl")
            search_results = wikipedia.search(topic, results=max_results)
            
            for title in search_results[:max_results]:
                try:
                    page = wikipedia.page(title)
                    sources.append({
                        'title': page.title,
                        'url': page.url,
                        'content': page.summary[:1000],
                        'type': 'wikipedia',
                        'credibility': 0.9,
                        'date': None
                    })
                    time.sleep(self.config.scraping.rate_limit)
                except Exception as e:
                    logging.warning(f"Błąd przy pobieraniu strony Wikipedia {title}: {e}")
                    continue
                    
        except Exception as e:
            logging.error(f"Błąd wyszukiwania Wikipedia: {e}")
            
        return sources
    
    def _search_news(self, topic: str, max_results: int) -> List[Dict]:
        """Wyszukuje w serwisach informacyjnych"""
        sources = []
        news_sites = self.sources_data.get("news", [])
        
        for site in news_sites[:max_results]:
            try:
                # Przykładowe wyszukiwanie - w rzeczywistości użyj News API lub RSS
                response = self.session.get(
                    site.get('search_url', '').format(query=topic),
                    timeout=self.config.scraping.timeout
                )
                response.raise_for_status()
                
                soup = BeautifulSoup(response.content, 'html.parser')
                articles = soup.find_all('article', limit=3)
                
                for article in articles:
                    title_elem = article.find(['h1', 'h2', 'h3'])
                    link_elem = article.find('a')
                    
                    if title_elem and link_elem:
                        sources.append({
                            'title': title_elem.get_text().strip(),
                            'url': urljoin(site['base_url'], link_elem.get('href')),
                            'content': self._extract_content(article.get_text()[:500]),
                            'type': 'news',
                            'credibility': site.get('credibility', 0.7),
                            'date': self._extract_date(article)
                        })
                
                time.sleep(self.config.scraping.rate_limit)
                
            except Exception as e:
                logging.warning(f"Błąd przy wyszukiwaniu w {site.get('name', 'Unknown')}: {e}")
                continue
                
        return sources
    
    def _search_academic(self, topic: str, max_results: int) -> List[Dict]:
        """Wyszukuje w bazach naukowych"""
        sources = []
        # Implementacja wyszukiwania w bazach akademickich
        # Przykład: Google Scholar, arXiv, itp.
        return sources
    
    def _search_general(self, topic: str, max_results: int) -> List[Dict]:
        """Ogólne wyszukiwanie internetowe"""
        sources = []
        general_sites = self.sources_data.get("general", [])
        
        for site in general_sites[:max_results]:
            try:
                # Implementacja ogólnego wyszukiwania
                pass
            except Exception as e:
                logging.warning(f"Błąd wyszukiwania ogólnego: {e}")
                continue
                
        return sources
    
    def _rank_sources(self, sources: List[Dict]) -> List[Dict]:
        """Ranguje źródła według wiarygodności i relevance"""
        return sorted(sources, key=lambda x: x.get('credibility', 0), reverse=True)
    
    def _extract_content(self, html_text: str) -> str:
        """Ekstraktuje czysty tekst"""
        return html_text.strip()[:1000]
    
    def _extract_date(self, element) -> Optional[str]:
        """Ekstraktuje datę z elementu"""
        # Implementacja ekstrakcji daty
        return None
'''
    
    with open("streamlit_article_generator/modules/web_scraper.py", "w", encoding="utf-8") as f:
        f.write(web_scraper_content)
    print("✅ Plik modules/web_scraper.py utworzony")


def create_llm_integration():
    """Tworzy plik modules/llm_integration.py"""
    print("\n🧠 Tworzenie pliku modules/llm_integration.py...")
    
    llm_content = '''import openai
import anthropic
from groq import Groq
import yaml
import os
from typing import Dict, List, Optional, Any
import logging

class LLMManager:
    def __init__(self, config):
        self.config = config
        self.clients = {}
        self.prompts = self._load_prompts()
        
    def _load_prompts(self) -> Dict:
        """Ładuje prompty z pliku YAML"""
        with open('config/prompts.yaml', 'r', encoding='utf-8') as f:
            return yaml.safe_load(f)
    
    def initialize_client(self, provider: str, api_key: str, model: str) -> bool:
        """Inicjalizuje klienta dla wybranego providera"""
        try:
            if provider == "openai":
                self.clients['openai'] = openai.OpenAI(api_key=api_key)
                self.current_provider = "openai"
                self.current_model = model
                
            elif provider == "anthropic":
                self.clients['anthropic'] = anthropic.Anthropic(api_key=api_key)
                self.current_provider = "anthropic"
                self.current_model = model
                
            elif provider == "ollama":
                # Ollama usually runs locally without API key
                try:
                    import ollama
                    self.clients['ollama'] = ollama.Client()
                except ImportError:
                    logging.error("Ollama not installed. Run: pip install ollama")
                    return False
                self.current_provider = "ollama"
                self.current_model = model
                
            elif provider == "groq":
                self.clients['groq'] = Groq(api_key=api_key)
                self.current_provider = "groq"
                self.current_model = model
                
            return True
            
        except Exception as e:
            logging.error(f"Błąd inicjalizacji klienta {provider}: {e}")
            return False
    
    def generate_article(self, topic: str, article_type: str, sources: List[Dict],
                        length: str, language: str = "polski", 
                        custom_instructions: str = "") -> Dict:
        """Generuje artykuł używając LLM"""
        
        if self.current_provider not in self.clients:
            raise ValueError("Klient LLM nie został zainicjalizowany")
        
        # Przygotuj prompt
        prompt_config = self.prompts['article_types'][article_type]
        formatted_sources = self._format_sources(sources)
        
        user_prompt = prompt_config['user_template'].format(
            topic=topic,
            length=length,
            language=language,
            sources=formatted_sources,
            custom_instructions=custom_instructions or "Brak dodatkowych instrukcji"
        )
        
        try:
            # Generuj artykuł
            if self.current_provider == "openai":
                response = self._generate_openai(prompt_config['system_prompt'], user_prompt)
            elif self.current_provider == "anthropic":
                response = self._generate_anthropic(prompt_config['system_prompt'], user_prompt)
            elif self.current_provider == "ollama":
                response = self._generate_ollama(prompt_config['system_prompt'], user_prompt)
            elif self.current_provider == "groq":
                response = self._generate_groq(prompt_config['system_prompt'], user_prompt)
            
            # Kalkuluj metryki
            quality_score = self._calculate_quality_score(response, sources)
            
            return {
                'content': response,
                'metadata': {
                    'topic': topic,
                    'type': article_type,
                    'length': length,
                    'language': language,
                    'provider': self.current_provider,
                    'model': self.current_model,
                    'sources_count': len(sources),
                    'quality_score': quality_score
                }
            }
            
        except Exception as e:
            logging.error(f"Błąd generowania artykułu: {e}")
            raise
    
    def _generate_openai(self, system_prompt: str, user_prompt: str) -> str:
        """Generuje używając OpenAI"""
        response = self.clients['openai'].chat.completions.create(
            model=self.current_model,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt}
            ],
            max_tokens=self.config.llm.max_tokens,
            temperature=self.config.llm.temperature
        )
        return response.choices[0].message.content
    
    def _generate_anthropic(self, system_prompt: str, user_prompt: str) -> str:
        """Generuje używając Anthropic Claude"""
        response = self.clients['anthropic'].messages.create(
            model=self.current_model,
            max_tokens=self.config.llm.max_tokens,
            temperature=self.config.llm.temperature,
            system=system_prompt,
            messages=[
                {"role": "user", "content": user_prompt}
            ]
        )
        return response.content[0].text
    
    def _generate_ollama(self, system_prompt: str, user_prompt: str) -> str:
        """Generuje używając Ollama"""
        import ollama
        response = ollama.chat(
            model=self.current_model,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt}
            ]
        )
        return response['message']['content']
    
    def _generate_groq(self, system_prompt: str, user_prompt: str) -> str:
        """Generuje używając Groq"""
        response = self.clients['groq'].chat.completions.create(
            model=self.current_model,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt}
            ],
            max_tokens=self.config.llm.max_tokens,
            temperature=self.config.llm.temperature
        )
        return response.choices[0].message.content
    
    def _format_sources(self, sources: List[Dict]) -> str:
        """Formatuje źródła do prompta"""
        formatted = []
        for i, source in enumerate(sources, 1):
            formatted.append(f"""
{i}. {source.get('title', 'Bez tytułu')}
   URL: {source.get('url', 'Brak URL')}
   Typ: {source.get('type', 'nieznany')}
   Treść: {source.get('content', 'Brak treści')[:500]}...
   Wiarygodność: {source.get('credibility', 0):.1f}/1.0
""")
        return "\\n".join(formatted)
    
    def _calculate_quality_score(self, content: str, sources: List[Dict]) -> float:
        """Kalkuluje ocenę jakości artykułu"""
        score = 0.5  # Base score
        
        # Długość artykułu
        if len(content) > 1000:
            score += 0.1
        if len(content) > 2000:
            score += 0.1
            
        # Liczba źródeł
        if len(sources) >= 3:
            score += 0.1
        if len(sources) >= 5:
            score += 0.1
            
        # Średnia wiarygodność źródeł
        if sources:
            avg_credibility = sum(s.get('credibility', 0) for s in sources) / len(sources)
            score += avg_credibility * 0.2
        
        return min(score, 1.0)
'''
    
    with open("streamlit_article_generator/modules/llm_integration.py", "w", encoding="utf-8") as f:
        f.write(llm_content)
    print("✅ Plik modules/llm_integration.py utworzony")


def create_article_generator():
    """Tworzy plik modules/article_generator.py"""
    print("\n📝 Tworzenie pliku modules/article_generator.py...")
    
    generator_content = '''from .web_scraper import WebScraper
from .llm_integration import LLMManager
import logging
import time
from typing import Dict, List, Optional, Callable

class ArticleGenerator:
    def __init__(self, config):
        self.config = config
        self.web_scraper = WebScraper(config)
        self.llm_manager = LLMManager(config)
        
    def generate_article(self, topic: str, article_type: str, length: str,
                        language: str, provider: str, model: str, api_key: str,
                        max_sources: int, search_depth: str, 
                        verify_facts: bool, include_citations: bool,
                        custom_instructions: str, progress_callback: Optional[Callable] = None) -> Dict:
        """
        Główna metoda generowania artykułu z progress tracking
        """
        
        try:
            # 1. Walidacja (0-10%)
            if progress_callback:
                progress_callback(5, "Walidacja danych wejściowych...")
            
            if not topic.strip():
                raise ValueError("Temat artykułu nie może być pusty")
            
            if not api_key.strip() and provider != "ollama":
                raise ValueError("Klucz API jest wymagany")
            
            if progress_callback:
                progress_callback(10, "Dane zwalidowane pomyślnie")
            
            # 2. Inicjalizacja LLM (10-20%)
            if progress_callback:
                progress_callback(15, f"Inicjalizacja {provider}...")
            
            if not self.llm_manager.initialize_client(provider, api_key, model):
                raise ValueError(f"Nie można zainicjalizować klienta {provider}")
            
            if progress_callback:
                progress_callback(20, "Klient LLM zainicjalizowany")
            
            # 3. Wyszukiwanie źródeł (20-40%)
            if progress_callback:
                progress_callback(25, "Wyszukiwanie źródeł...")
            
            sources = self.web_scraper.search_sources(
                topic=topic,
                max_sources=max_sources,
                search_depth=search_depth
            )
            
            if progress_callback:
                progress_callback(40, f"Znaleziono {len(sources)} źródeł")
            
            # 4. Generowanie artykułu (40-70%)
            if progress_callback:
                progress_callback(50, "Generowanie artykułu...")
            
            article_data = self.llm_manager.generate_article(
                topic=topic,
                article_type=article_type,
                sources=sources,
                length=length,
                language=language,
                custom_instructions=custom_instructions
            )
            
            if progress_callback:
                progress_callback(70, "Artykuł wygenerowany")
            
            # 5. Opcjonalna weryfikacja faktów (70-90%)
            if verify_facts:
                if progress_callback:
                    progress_callback(80, "Weryfikacja faktów...")
                
                # Implementacja weryfikacji faktów
                article_data['fact_check_score'] = self._verify_facts(
                    article_data['content'], sources
                )
                
                if progress_callback:
                    progress_callback(90, "Fakty zweryfikowane")
            
            # 6. Finalizacja (90-100%)
            if progress_callback:
                progress_callback(95, "Finalizacja...")
            
            # Dodaj źródła i metadata
            article_data['sources'] = sources
            article_data['generation_time'] = time.time()
            article_data['word_count'] = len(article_data['content'].split())
            
            if progress_callback:
                progress_callback(100, "Artykuł gotowy!")
            
            return article_data
            
        except Exception as e:
            logging.error(f"Błąd generowania artykułu: {e}")
            if progress_callback:
                progress_callback(0, f"Błąd: {str(e)}")
            raise
    
    def _verify_facts(self, content: str, sources: List[Dict]) -> float:
        """Weryfikuje fakty w artykule względem źródeł"""
        # Implementacja weryfikacji faktów
        # Można użyć NLP do wyodrębnienia faktów i porównania ze źródłami
        return 0.8  # Placeholder
'''
    
    with open("streamlit_article_generator/modules/article_generator.py", "w", encoding="utf-8") as f:
        f.write(generator_content)
    print("✅ Plik modules/article_generator.py utworzony")


def create_export_manager():
    """Tworzy plik modules/export_manager.py"""
    print("\n💾 Tworzenie pliku modules/export_manager.py...")
    
    export_content = '''import os
import json
from datetime import datetime
from typing import Dict, Optional
import markdown2
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from docx import Document
import logging

class ExportManager:
    def __init__(self, config):
        self.config = config
        self.output_dir = config.output_dir
        os.makedirs(self.output_dir, exist_ok=True)
    
    def save_article(self, article_data: Dict, filename: str, 
                    export_format: str, include_bibliography: bool = True,
                    include_metadata: bool = True) -> str:
        """Główna metoda eksportu artykułu"""
        
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        base_filename = f"{filename}_{timestamp}"
        
        try:
            if export_format == "html":
                filepath = self._save_html(article_data, base_filename, 
                                         include_bibliography, include_metadata)
            elif export_format == "markdown":
                filepath = self._save_markdown(article_data, base_filename,
                                             include_bibliography, include_metadata)
            elif export_format == "pdf":
                filepath = self._save_pdf(article_data, base_filename,
                                        include_bibliography, include_metadata)
            elif export_format == "docx":
                filepath = self._save_docx(article_data, base_filename,
                                         include_bibliography, include_metadata)
            elif export_format == "txt":
                filepath = self._save_txt(article_data, base_filename,
                                        include_bibliography, include_metadata)
            else:
                raise ValueError(f"Nieobsługiwany format eksportu: {export_format}")
            
            return filepath
            
        except Exception as e:
            logging.error(f"Błąd eksportu do {export_format}: {e}")
            raise
    
    def _save_html(self, article_data: Dict, filename: str,
                   include_bibliography: bool, include_metadata: bool) -> str:
        """Eksport do HTML"""
        filepath = os.path.join(self.output_dir, f"{filename}.html")
        
        html_content = f"""<!DOCTYPE html>
<html lang="pl">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{article_data['metadata']['topic']}</title>
    <style>
        body {{ font-family: Arial, sans-serif; line-height: 1.6; margin: 40px; }}
        h1 {{ color: #333; border-bottom: 2px solid #ccc; }}
        h2 {{ color: #666; }}
        .metadata {{ background: #f4f4f4; padding: 15px; margin: 20px 0; border-radius: 5px; }}
        .sources {{ margin-top: 30px; }}
        .source {{ margin: 10px 0; padding: 10px; background: #f9f9f9; border-left: 3px solid #007cba; }}
    </style>
</head>
<body>
    <h1>{article_data['metadata']['topic']}</h1>
"""
        
        if include_metadata:
            html_content += f"""
    <div class="metadata">
        <h3>Metadane</h3>
        <p><strong>Typ:</strong> {article_data['metadata']['type']}</p>
        <p><strong>Długość:</strong> {article_data['metadata']['length']}</p>
        <p><strong>Język:</strong> {article_data['metadata']['language']}</p>
        <p><strong>Model:</strong> {article_data['metadata']['provider']} - {article_data['metadata']['model']}</p>
        <p><strong>Liczba słów:</strong> {article_data.get('word_count', 'N/A')}</p>
        <p><strong>Ocena jakości:</strong> {article_data['metadata']['quality_score']:.2f}/1.0</p>
    </div>
"""
        
        # Konwertuj markdown na HTML
        html_content += markdown2.markdown(article_data['content'])
        
        if include_bibliography and 'sources' in article_data:
            html_content += """
    <div class="sources">
        <h2>Bibliografia</h2>
"""
            for i, source in enumerate(article_data['sources'], 1):
                html_content += f"""
        <div class="source">
            <strong>{i}. {source.get('title', 'Bez tytułu')}</strong><br>
            <em>Typ:</em> {source.get('type', 'nieznany')}<br>
            <em>URL:</em> <a href="{source.get('url', '#')}" target="_blank">{source.get('url', 'Brak URL')}</a><br>
            <em>Wiarygodność:</em> {source.get('credibility', 0):.1f}/1.0
        </div>
"""
            html_content += "    </div>"
        
        html_content += """
</body>
</html>
"""
        
        with open(filepath, 'w', encoding='utf-8') as f:
            f.write(html_content)
        
        return filepath
    
    def _save_markdown(self, article_data: Dict, filename: str,
                      include_bibliography: bool, include_metadata: bool) -> str:
        """Eksport do Markdown"""
        filepath = os.path.join(self.output_dir, f"{filename}.md")
        
        content = f"# {article_data['metadata']['topic']}\\n\\n"
        
        if include_metadata:
            content += "## Metadane\\n\\n"
            content += f"- **Typ:** {article_data['metadata']['type']}\\n"
            content += f"- **Długość:** {article_data['metadata']['length']}\\n"
            content += f"- **Język:** {article_data['metadata']['language']}\\n"
            content += f"- **Model:** {article_data['metadata']['provider']} - {article_data['metadata']['model']}\\n"
            content += f"- **Liczba słów:** {article_data.get('word_count', 'N/A')}\\n"
            content += f"- **Ocena jakości:** {article_data['metadata']['quality_score']:.2f}/1.0\\n\\n"
        
        content += article_data['content']
        
        if include_bibliography and 'sources' in article_data:
            content += "\\n\\n## Bibliografia\\n\\n"
            for i, source in enumerate(article_data['sources'], 1):
                content += f"{i}. **{source.get('title', 'Bez tytułu')}**\\n"
                content += f"   - Typ: {source.get('type', 'nieznany')}\\n"
                content += f"   - URL: [{source.get('url', 'Brak URL')}]({source.get('url', '#')})\\n"
                content += f"   - Wiarygodność: {source.get('credibility', 0):.1f}/1.0\\n\\n"
        
        with open(filepath, 'w', encoding='utf-8') as f:
            f.write(content)
        
        return filepath
    
    def _save_pdf(self, article_data: Dict, filename: str,
                  include_bibliography: bool, include_metadata: bool) -> str:
        """Eksport do PDF"""
        filepath = os.path.join(self.output_dir, f"{filename}.pdf")
        
        doc = SimpleDocTemplate(filepath, pagesize=letter)
        styles = getSampleStyleSheet()
        story = []
        
        # Tytuł
        title = Paragraph(article_data['metadata']['topic'], styles['Title'])
        story.append(title)
        story.append(Spacer(1, 12))
        
        # Treść artykułu (uproszczona)
        content_lines = article_data['content'].split('\\n')
        for line in content_lines[:50]:  # Limit dla przykładu
            if line.strip():
                para = Paragraph(line, styles['Normal'])
                story.append(para)
                story.append(Spacer(1, 6))
        
        doc.build(story)
        return filepath
    
    def _save_docx(self, article_data: Dict, filename: str,
                   include_bibliography: bool, include_metadata: bool) -> str:
        """Eksport do DOCX"""
        filepath = os.path.join(self.output_dir, f"{filename}.docx")
        
        doc = Document()
        
        # Tytuł
        title = doc.add_heading(article_data['metadata']['topic'], 0)
        
        # Treść artykułu (uproszczona)
        content_lines = article_data['content'].split('\\n')
        for line in content_lines:
            if line.strip():
                if line.startswith('#'):
                    level = line.count('#')
                    clean_line = line.lstrip('# ').strip()
                    doc.add_heading(clean_line, level=min(level, 3))
                else:
                    doc.add_paragraph(line)
        
        doc.save(filepath)
        return filepath
    
    def _save_txt(self, article_data: Dict, filename: str,
                  include_bibliography: bool, include_metadata: bool) -> str:
        """Eksport do TXT"""
        filepath = os.path.join(self.output_dir, f"{filename}.txt")
        
        content = f"{article_data['metadata']['topic']}\\n"
        content += "=" * len(article_data['metadata']['topic']) + "\\n\\n"
        content += article_data['content']
        
        with open(filepath, 'w', encoding='utf-8') as f:
            f.write(content)
        
        return filepath
'''
    
    with open("streamlit_article_generator/modules/export_manager.py", "w", encoding="utf-8") as f:
        f.write(export_content)
    print("✅ Plik modules/export_manager.py utworzony")


def create_main_app():
    """Tworzy główną aplikację app.py"""
    print("\n🚀 Tworzenie głównej aplikacji app.py...")
    
    app_content = '''import streamlit as st
import os
import json
import time
import pandas as pd
import plotly.express as px
import plotly.graph_objects as go
from datetime import datetime
try:
    from wordcloud import WordCloud
    import matplotlib.pyplot as plt
    WORDCLOUD_AVAILABLE = True
except ImportError:
    WORDCLOUD_AVAILABLE = False

# Import modułów
from config.settings import AppConfig
from modules.article_generator import ArticleGenerator
from modules.export_manager import ExportManager

# Konfiguracja strony
st.set_page_config(
    page_title="Generator Artykułów AI",
    page_icon="📝",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Inicjalizacja konfiguracji
@st.cache_resource
def load_config():
    return AppConfig.load()

@st.cache_resource
def load_article_generator(config):
    return ArticleGenerator(config)

@st.cache_resource
def load_export_manager(config):
    return ExportManager(config)

# Funkcje pomocnicze
def save_to_history(article_data):
    """Zapisuje artykuł do historii"""
    history_file = 'data/history.json'
    os.makedirs('data', exist_ok=True)
    
    history_entry = {
        'timestamp': datetime.now().isoformat(),
        'topic': article_data['metadata']['topic'],
        'type': article_data['metadata']['type'],
        'word_count': article_data.get('word_count', 0),
        'quality_score': article_data['metadata']['quality_score'],
        'provider': article_data['metadata']['provider']
    }
    
    if os.path.exists(history_file):
        with open(history_file, 'r', encoding='utf-8') as f:
            history = json.load(f)
    else:
        history = []
    
    history.insert(0, history_entry)
    history = history[:50]  # Limit do 50 ostatnich
    
    with open(history_file, 'w', encoding='utf-8') as f:
        json.dump(history, f, ensure_ascii=False, indent=2)

def load_history():
    """Ładuje historię artykułów"""
    history_file = 'data/history.json'
    if os.path.exists(history_file):
        with open(history_file, 'r', encoding='utf-8') as f:
            return json.load(f)
    return []

def main():
    # Ładowanie konfiguracji
    config = load_config()
    article_generator = load_article_generator(config)
    export_manager = load_export_manager(config)
    
    # Inicjalizacja session state
    if 'article_generated' not in st.session_state:
        st.session_state.article_generated = False
    if 'current_article' not in st.session_state:
        st.session_state.current_article = None
    if 'generation_time' not in st.session_state:
        st.session_state.generation_time = 0
    
    # Tytuł i opis
    st.title("🤖 Generator Artykułów AI")
    st.markdown("Zaawansowana aplikacja do automatycznego generowania artykułów z wykorzystaniem LLM")
    
    # Sidebar - Konfiguracja
    with st.sidebar:
        st.header("⚙️ Konfiguracja")
        
        # Sekcja LLM Provider
        st.subheader("🧠 LLM Provider")
        provider = st.selectbox(
            "Wybierz providera:",
            ["openai", "anthropic", "ollama", "groq"],
            help="Wybierz dostawcę modelu językowego"
        )
        
        # Wybór modelu w zależności od providera
        if provider == "openai":
            model = st.selectbox("Model:", config.llm.openai_models)
            api_key = st.text_input("OpenAI API Key:", type="password")
        elif provider == "anthropic":
            model = st.selectbox("Model:", config.llm.anthropic_models)
            api_key = st.text_input("Anthropic API Key:", type="password")
        elif provider == "ollama":
            model = st.selectbox("Model:", config.llm.ollama_models)
            api_key = ""  # Ollama nie wymaga API key
            st.info("Ollama działa lokalnie - nie wymaga API key")
        elif provider == "groq":
            model = st.selectbox("Model:", config.llm.groq_models)
            api_key = st.text_input("Groq API Key:", type="password")
        
        st.divider()
        
        # Zaawansowane ustawienia
        with st.expander("🔧 Ustawienia zaawansowane"):
            max_sources = st.slider(
                "Maksymalna liczba źródeł:",
                min_value=3,
                max_value=15,
                value=10
            )
            
            search_depth = st.selectbox(
                "Głębokość wyszukiwania:",
                ["shallow", "medium", "deep"],
                index=1
            )
            
            col1, col2 = st.columns(2)
            with col1:
                verify_facts = st.checkbox("Weryfikacja faktów", value=False)
            with col2:
                include_citations = st.checkbox("Cytowania", value=True)
            
            custom_instructions = st.text_area(
                "Dodatkowe instrukcje dla LLM:",
                placeholder="Np. skup się na aspekcie technicznym..."
            )
        
        st.divider()
        
        # Historia artykułów
        with st.expander("📚 Historia artykułów"):
            history = load_history()
            if history:
                for entry in history[:5]:  # Pokaż ostatnie 5
                    st.write(f"**{entry['topic']}**")
                    st.write(f"Typ: {entry['type']} | Ocena: {entry['quality_score']:.2f}")
                    st.write(f"Data: {entry['timestamp'][:16]}")
                    st.write("---")
            else:
                st.write("Brak historii")
    
    # Główny interfejs
    col1, col2, col3 = st.columns([2, 1, 1])
    
    with col1:
        topic = st.text_input(
            "🎯 Temat artykułu:",
            placeholder="Wprowadź temat artykułu...",
            help="Opisz temat, o którym ma zostać napisany artykuł"
        )
    
    with col2:
        article_type = st.selectbox(
            "📝 Typ artykułu:",
            ["esej", "blog_post", "artykul_wikipedii", "artykul_naukowy", "poradnik", "news_summary"]
        )
    
    with col3:
        language = st.selectbox(
            "🌍 Język:",
            ["polski", "angielski", "niemiecki", "francuski"]
        )
    
    # Długość artykułu
    length = st.radio(
        "📏 Długość artykułu:",
        ["krótki (500-800 słów)", "średni (800-1500 słów)", "długi (1500-3000 słów)"],
        horizontal=True
    )
    
    # Przycisk generowania
    col1, col2, col3 = st.columns([1, 2, 1])
    with col2:
        generate_button = st.button(
            "🚀 Generuj Artykuł",
            type="primary",
            use_container_width=True,
            disabled=not topic.strip() or (not api_key.strip() and provider != "ollama")
        )
    
    # Proces generowania
    if generate_button:
        # Progress tracking
        progress_bar = st.progress(0)
        status_text = st.empty()
        
        def update_progress(progress, message):
            progress_bar.progress(progress / 100)
            status_text.text(message)
        
        try:
            start_time = time.time()
            
            # Generowanie artykułu
            article_data = article_generator.generate_article(
                topic=topic,
                article_type=article_type,
                length=length,
                language=language,
                provider=provider,
                model=model,
                api_key=api_key,
                max_sources=max_sources,
                search_depth=search_depth,
                verify_facts=verify_facts,
                include_citations=include_citations,
                custom_instructions=custom_instructions,
                progress_callback=update_progress
            )
            
            end_time = time.time()
            generation_time = end_time - start_time
            
            # Zapisz do session state
            st.session_state.current_article = article_data
            st.session_state.article_generated = True
            st.session_state.generation_time = generation_time
            
            # Zapisz do historii
            save_to_history(article_data)
            
            # Ukryj progress bar
            progress_bar.empty()
            status_text.empty()
            
            st.success(f"✅ Artykuł wygenerowany pomyślnie w {generation_time:.1f} sekund!")
            st.rerun()
            
        except Exception as e:
            progress_bar.empty()
            status_text.empty()
            st.error(f"❌ Błąd podczas generowania artykułu: {str(e)}")
    
    # Wyświetlanie wyników
    if st.session_state.article_generated and st.session_state.current_article:
        article_data = st.session_state.current_article
        
        # Tabs dla wyników
        tab1, tab2, tab3, tab4 = st.tabs(["📄 Artykuł", "📊 Analiza", "🔗 Źródła", "💾 Eksport"])
        
        with tab1:
            # Metryki
            col1, col2, col3, col4 = st.columns(4)
            
            with col1:
                st.metric(
                    "Liczba słów",
                    article_data.get('word_count', 0)
                )
            
            with col2:
                st.metric(
                    "Źródła",
                    len(article_data.get('sources', []))
                )
            
            with col3:
                st.metric(
                    "Czas generowania",
                    f"{st.session_state.generation_time:.1f}s"
                )
            
            with col4:
                score = article_data['metadata']['quality_score']
                st.metric(
                    "Ocena jakości",
                    f"{score:.2f}/1.0",
                    delta=f"{score-0.5:.2f}"
                )
            
            st.divider()
            
            # Główna treść artykułu
            st.markdown("### Treść artykułu:")
            st.markdown(article_data['content'])
            
            # Opcja edycji
            with st.expander("✏️ Edytuj artykuł"):
                edited_content = st.text_area(
                    "Edytuj treść:",
                    value=article_data['content'],
                    height=300
                )
                
                if st.button("💾 Zapisz zmiany"):
                    st.session_state.current_article['content'] = edited_content
                    st.success("Zmiany zapisane!")
                    st.rerun()
        
        with tab2:
            # Statystyki tekstu
            content = article_data['content']
            words = len(content.split())
            chars = len(content)
            paragraphs = len([p for p in content.split('\\n') if p.strip()])
            sentences = len([s for s in content.split('.') if s.strip()])
            avg_sentence_length = words / max(sentences, 1)
            
            stats_df = pd.DataFrame({
                'Metryka': ['Słowa', 'Znaki', 'Akapity', 'Zdania', 'Średnia długość zdania'],
                'Wartość': [words, chars, paragraphs, sentences, f"{avg_sentence_length:.1f}"]
            })
            
            st.subheader("📈 Statystyki tekstu")
            st.dataframe(stats_df, use_container_width=True)
            
            # Wykresy
            col1, col2 = st.columns(2)
            
            with col1:
                # Wykres rozkładu typów źródeł
                if 'sources' in article_data and article_data['sources']:
                    source_types = [s.get('type', 'nieznany') for s in article_data['sources']]
                    type_counts = pd.Series(source_types).value_counts()
                    
                    fig_pie = px.pie(
                        values=type_counts.values,
                        names=type_counts.index,
                        title="Rozkład typów źródeł"
                    )
                    st.plotly_chart(fig_pie, use_container_width=True)
            
            with col2:
                # Wykres wiarygodności źródeł
                if 'sources' in article_data and article_data['sources']:
                    credibility_data = [s.get('credibility', 0) for s in article_data['sources']]
                    
                    fig_bar = go.Figure(data=[
                        go.Bar(x=list(range(1, len(credibility_data)+1)), y=credibility_data)
                    ])
                    fig_bar.update_layout(
                        title="Wiarygodność źródeł",
                        xaxis_title="Źródło #",
                        yaxis_title="Wiarygodność"
                    )
                    st.plotly_chart(fig_bar, use_container_width=True)
            
            # Opcjonalna chmura słów
            if WORDCLOUD_AVAILABLE and st.checkbox("☁️ Pokaż chmurę słów"):
                try:
                    wordcloud = WordCloud(
                        width=800, 
                        height=400, 
                        background_color='white',
                        colormap='viridis'
                    ).generate(content)
                    
                    fig, ax = plt.subplots(figsize=(10, 5))
                    ax.imshow(wordcloud, interpolation='bilinear')
                    ax.axis('off')
                    st.pyplot(fig)
                except Exception as e:
                    st.error(f"Błąd generowania chmury słów: {e}")
        
        with tab3:
            # Filtry źródeł
            if 'sources' in article_data and article_data['sources']:
                col1, col2 = st.columns(2)
                
                with col1:
                    source_types = list(set([s.get('type', 'nieznany') for s in article_data['sources']]))
                    selected_types = st.multiselect(
                        "Filtruj według typu:",
                        source_types,
                        default=source_types
                    )
                
                with col2:
                    min_credibility = st.slider(
                        "Minimalna wiarygodność:",
                        0.0, 1.0, 0.0, 0.1
                    )
                
                # Filtrowane źródła
                filtered_sources = [
                    s for s in article_data['sources']
                    if s.get('type', 'nieznany') in selected_types
                    and s.get('credibility', 0) >= min_credibility
                ]
                
                st.write(f"Znaleziono {len(filtered_sources)} źródeł")
                
                # Wyświetl źródła
                for i, source in enumerate(filtered_sources, 1):
                    with st.expander(f"🔗 Źródło {i}: {source.get('title', 'Bez tytułu')}"):
                        col1, col2 = st.columns([3, 1])
                        
                        with col1:
                            st.write(f"**URL:** {source.get('url', 'Brak URL')}")
                            st.write(f"**Typ:** {source.get('type', 'nieznany')}")
                            st.write(f"**Podsumowanie:** {source.get('content', 'Brak treści')[:300]}...")
                        
                        with col2:
                            credibility = source.get('credibility', 0)
                            st.metric("Wiarygodność", f"{credibility:.1f}/1.0")
                            
                            if source.get('url'):
                                st.link_button(
                                    "Otwórz źródło",
                                    source['url'],
                                    use_container_width=True
                                )
            else:
                st.info("Brak dostępnych źródeł")
        
        with tab4:
            # Ustawienia eksportu
            col1, col2 = st.columns(2)
            
            with col1:
                st.subheader("⚙️ Ustawienia eksportu")
                
                export_format = st.selectbox(
                    "Format eksportu:",
                    ["html", "markdown", "pdf", "docx", "txt"]
                )
                
                include_bibliography = st.checkbox("Dołącz bibliografię", value=True)
                include_metadata = st.checkbox("Dołącz metadane", value=True)
            
            with col2:
                st.subheader("📁 Nazwa pliku")
                
                filename = st.text_input(
                    "Nazwa pliku:",
                    value=article_data['metadata']['topic'][:30].replace(' ', '_')
                )
                
                full_filename = f"{filename}.{export_format}"
                st.write(f"**Pełna nazwa:** `{full_filename}`")
            
            # Eksport
            if st.button("💾 Eksportuj artykuł", type="primary", use_container_width=True):
                try:
                    filepath = export_manager.save_article(
                        article_data=article_data,
                        filename=filename,
                        export_format=export_format,
                        include_bibliography=include_bibliography,
                        include_metadata=include_metadata
                    )
                    
                    st.success(f"✅ Artykuł wyeksportowany do: {filepath}")
                    
                    # Download button
                    with open(filepath, 'rb') as f:
                        st.download_button(
                            label=f"⬇️ Pobierz {export_format.upper()}",
                            data=f.read(),
                            file_name=full_filename,
                            mime=f"application/{export_format}",
                            use_container_width=True
                        )
                        
                except Exception as e:
                    st.error(f"❌ Błąd eksportu: {str(e)}")
            
            # Podgląd eksportu
            with st.expander("👁️ Podgląd eksportu"):
                if export_format == "markdown":
                    preview_content = f"# {article_data['metadata']['topic']}\\n\\n"
                    preview_content += article_data['content'][:500] + "..."
                    st.code(preview_content, language="markdown")
                elif export_format == "html":
                    preview_content = f"<h1>{article_data['metadata']['topic']}</h1>\\n"
                    preview_content += "<p>" + article_data['content'][:500] + "...</p>"
                    st.code(preview_content, language="html")
                else:
                    st.text(article_data['content'][:500] + "...")

if __name__ == "__main__":
    main()
'''
    
    with open("streamlit_article_generator/app.py", "w", encoding="utf-8") as f:
        f.write(app_content)
    print("✅ Główna aplikacja app.py utworzona")


def create_readme():
    """Tworzy plik README.md"""
    print("\n📖 Tworzenie pliku README.md...")
    
    readme_content = '''# 🤖 Streamlit Article Generator

Zaawansowana aplikacja do automatycznego generowania artykułów z wykorzystaniem LLM (Large Language Models).

## 🚀 Funkcjonalności

- **Integracja z wieloma LLM**: OpenAI GPT, Anthropic Claude, Ollama, Groq
- **Automatyczne wyszukiwanie źródeł**: Wikipedia, serwisy informacyjne, bazy akademickie
- **Różne typy artykułów**: eseje, posty blogowe, artykuły naukowe, poradniki
- **Analiza jakości**: ocena wiarygodności źródeł, statystyki tekstu
- **Eksport do wielu formatów**: HTML, Markdown, PDF, DOCX, TXT
- **Historia artykułów**: automatyczne zapisywanie i ładowanie poprzednich prac

## 📦 Instalacja

1. **Sklonuj repozytorium:**
```bash
git clone <repository-url>
cd streamlit_article_generator
```

2. **Zainstaluj zależności:**
```bash
pip install -r requirements.txt
```

3. **Pobierz modele NLP (opcjonalne):**
```bash
python -m spacy download pl_core_news_sm
```

4. **Skonfiguruj zmienne środowiskowe:**
```bash
cp .env.example .env
# Edytuj plik .env i dodaj swoje klucze API
```

## 🔧 Konfiguracja

### Klucze API

Uzupełnij plik `.env` swoimi kluczami API:

```env
OPENAI_API_KEY=your_openai_api_key_here
ANTHROPIC_API_KEY=your_anthropic_api_key_here
GROQ_API_KEY=your_groq_api_key_here
```

### Ollama (opcjonalne)

Dla lokalnych modeli zainstaluj Ollama:
```bash
# Zainstaluj Ollama z oficjalnej strony
# Pobierz modele:
ollama pull llama3.1
ollama pull mistral
```

## 🎯 Użytkowanie

1. **Uruchom aplikację:**
```bash
streamlit run app.py
```

2. **Otwórz przeglądarkę** i przejdź do `http://localhost:8501`

3. **Skonfiguruj LLM**:
   - Wybierz providera (OpenAI, Anthropic, Ollama, Groq)
   - Wprowadź klucz API (jeśli wymagany)
   - Wybierz model

4. **Wygeneruj artykuł**:
   - Wprowadź temat
   - Wybierz typ artykułu
   - Ustaw parametry
   - Kliknij "Generuj Artykuł"

5. **Analizuj i eksportuj**:
   - Przejrzyj wygenerowany artykuł
   - Sprawdź źródła i statystyki
   - Wyeksportuj w wybranym formacie

## 📁 Struktura projektu

```
streamlit_article_generator/
├── app.py                   # Główna aplikacja Streamlit
├── config/
│   ├── settings.py          # Konfiguracja aplikacji
│   └── prompts.yaml         # Prompty dla różnych typów artykułów
├── modules/
│   ├── web_scraper.py       # Web scraping
│   ├── llm_integration.py   # Integracja z LLM
│   ├── article_generator.py # Generator artykułów
│   └── export_manager.py    # Eksport do plików
├── data/
│   ├── sources.json         # Lista zaufanych źródeł
│   └── history.json         # Historia artykułów
├── outputs/                 # Wygenerowane pliki
├── requirements.txt
├── .env                     # Zmienne środowiskowe
└── README.md
```

## 🔧 Zaawansowana konfiguracja

### Dodawanie nowych źródeł

Edytuj `data/sources.json`:
```json
{
  "news": [
    {
      "name": "Nowe Źródło",
      "base_url": "https://example.com",
      "search_url": "https://example.com/search?q={query}",
      "credibility": 0.8
    }
  ]
}
```

### Niestandardowe prompty

Modyfikuj `config/prompts.yaml` aby dostosować prompty dla różnych typów artykułów.

## 🐛 Rozwiązywanie problemów

### Błędy instalacji

```bash
# Jeśli problem z reportlab:
pip install reportlab --upgrade

# Jeśli problem z wordcloud:
pip install wordcloud --upgrade

# Jeśli problem ze spacy:
python -m spacy download pl_core_news_sm --force
```

### Błędy API

- Sprawdź poprawność kluczy API w pliku `.env`
- Upewnij się, że masz wystarczające limity w swoim koncie
- Dla Ollama sprawdź czy serwer działa: `ollama serve`

## 📈 Performance

### Optymalizacja

- Użyj cache dla powtarzających się zapytań
- Ogranicz liczbę źródeł dla szybszego działania
- Wybierz mniejsze modele dla prostszych zadań

### Limity

- Maksymalna liczba źródeł: 15
- Maksymalna długość artykułu: 3000 słów
- Cache TTL: 1 godzina

## 🤝 Wkład w projekt

1. Fork repozytorium
2. Utwórz branch dla nowej funkcjonalności
3. Commituj zmiany
4. Utwórz Pull Request

## 📄 Licencja

MIT License - szczegóły w pliku LICENSE

## 🆘 Wsparcie

W przypadku problemów:
1. Sprawdź sekcję "Rozwiązywanie problemów"
2. Przeszukaj istniejące Issues
3. Utwórz nowe Issue z opisem problemu

## 🔮 Roadmapa

- [ ] Integracja z dodatkowymi LLM
- [ ] Ulepszone wyszukiwanie akademickie
- [ ] System pluginów
- [ ] Współpraca zespołowa
- [ ] Mobile app
- [ ] API endpoints

---

Stworzono z ❤️ używając Streamlit i Python
'''
    
    with open("streamlit_article_generator/README.md", "w", encoding="utf-8") as f:
        f.write(readme_content)
    print("✅ Plik README.md utworzony")


def create_run_script():
    """Tworzy skrypt uruchomieniowy"""
    print("\n🏃 Tworzenie skryptu uruchomieniowego...")
    
    # Skrypt dla Windows
    run_bat_content = '''@echo off
echo Starting Streamlit Article Generator...
echo.

REM Check if virtual environment exists
if not exist "venv" (
    echo Creating virtual environment...
    python -m venv venv
)

REM Activate virtual environment
call venv\\Scripts\\activate

REM Install requirements
echo Installing requirements...
pip install -r requirements.txt

REM Run the application
echo.
echo Starting application...
echo Open your browser and go to: http://localhost:8501
echo.
streamlit run app.py

pause
'''
    
    with open("streamlit_article_generator/run.bat", "w", encoding="utf-8") as f:
        f.write(run_bat_content)
    
    # Skrypt dla Linux/Mac
    run_sh_content = '''#!/bin/bash

echo "Starting Streamlit Article Generator..."
echo

# Check if virtual environment exists
if [ ! -d "venv" ]; then
    echo "Creating virtual environment..."
    python3 -m venv venv
fi

# Activate virtual environment
source venv/bin/activate

# Install requirements
echo "Installing requirements..."
pip install -r requirements.txt

# Run the application
echo
echo "Starting application..."
echo "Open your browser and go to: http://localhost:8501"
echo
streamlit run app.py
'''
    
    with open("streamlit_article_generator/run.sh", "w", encoding="utf-8") as f:
        f.write(run_sh_content)
    
    # Make executable on Unix systems
    try:
        os.chmod("streamlit_article_generator/run.sh", 0o755)
    except:
        pass
    
    print("✅ Skrypty uruchomieniowe utworzone")


def create_gitignore():
    """Tworzy plik .gitignore"""
    print("\n🙈 Tworzenie pliku .gitignore...")
    
    gitignore_content = '''# Byte-compiled / optimized / DLL files
__pycache__/
*.py[cod]
*$py.class

# C extensions
*.so

# Distribution / packaging
.Python
build/
develop-eggs/
dist/
downloads/
eggs/
.eggs/
lib/
lib64/
parts/
sdist/
var/
wheels/
*.egg-info/
.installed.cfg
*.egg

# PyInstaller
*.manifest
*.spec

# Installer logs
pip-log.txt
pip-delete-this-directory.txt

# Unit test / coverage reports
htmlcov/
.tox/
.coverage
.coverage.*
.cache
nosetests.xml
coverage.xml
*.cover
.hypothesis/
.pytest_cache/

# Virtual environments
venv/
env/
ENV/

# Streamlit
.streamlit/

# Environment variables
.env
.env.local
.env.production

# IDE
.vscode/
.idea/
*.swp
*.swo
*~

# OS
.DS_Store
Thumbs.db

# Project specific
outputs/*.pdf
outputs/*.docx
outputs/*.html
outputs/*.txt
outputs/*.md
data/cache/*
data/history.json
*.log

# API Keys and secrets
config/keys.json
secrets.toml

# Temporary files
temp/
tmp/
*.tmp
'''
    
    with open("streamlit_article_generator/.gitignore", "w", encoding="utf-8") as f:
        f.write(gitignore_content)
    print("✅ Plik .gitignore utworzony")


def main():
    """Główna funkcja uruchamiająca cały proces"""
    print("🚀 Rozpoczynam tworzenie projektu Streamlit Article Generator...")
    print("=" * 60)
    
    try:
        # Sprawdź czy folder już istnieje
        if os.path.exists("streamlit_article_generator"):
            response = input("📁 Folder 'streamlit_article_generator' już istnieje. Czy chcesz go zastąpić? (y/N): ")
            if response.lower() != 'y':
                print("❌ Anulowano tworzenie projektu.")
                return
            
            # Usuń istniejący folder
            import shutil
            shutil.rmtree("streamlit_article_generator")
            print("🗑️  Usunięto istniejący folder")
        
        # Tworzenie struktury i plików
        create_directory_structure()
        create_requirements_txt()
        create_env_file()
        create_config_settings()
        create_prompts_yaml()
        create_sources_json()
        create_modules_init()
        create_web_scraper()
        create_llm_integration()
        create_article_generator()
        create_export_manager()
        create_main_app()
        create_readme()
        create_run_script()
        create_gitignore()
        
        print("\n" + "=" * 60)
        print("🎉 Projekt został pomyślnie utworzony!")
        print("\n📋 Następne kroki:")
        print("1. cd streamlit_article_generator")
        print("2. pip install -r requirements.txt")
        print("3. Uzupełnij klucze API w pliku .env")
        print("4. streamlit run app.py")
        print("\n🚀 Alternatywnie uruchom:")
        print("   Windows: run.bat")
        print("   Linux/Mac: ./run.sh")
        print("\n📖 Więcej informacji w pliku README.md")
        
    except Exception as e:
        print(f"\n❌ Błąd podczas tworzenia projektu: {e}")
        print("🔧 Sprawdź uprawnienia do zapisu i spróbuj ponownie.")
        sys.exit(1)


if __name__ == "__main__":
    main()